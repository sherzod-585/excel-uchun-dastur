import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
import pandas as pd

MONTHS_RU = {
    'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
    'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
    'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12
}

MONTHS_UZ = {
    'yanvar': 1, 'fevral': 2, 'mart': 3, 'aprel': 4,
    'may': 5, 'iyun': 6, 'iyul': 7, 'avgust': 8,
    'sentabr': 9, 'oktabr': 10, 'noyabr': 11, 'dekabr': 12
}


UZB_CYRILLIC_MAP = {
    'қ': 'к', 'Қ': 'К',
    'ғ': 'г', 'Ғ': 'Г',
    'ў': 'у', 'Ў': 'У',
    'ҳ': 'х', 'Ҳ': 'Х',
}

LATIN_MULTI = [
    ("O'", 'У'), ("o'", 'у'), ("O`", 'У'), ("o`", 'у'),
    ("G'", 'Г'), ("g'", 'г'), ("G`", 'Г'), ("g`", 'г'),
    ('SH', 'Ш'), ('Sh', 'Ш'), ('sh', 'ш'),
    ('CH', 'Ч'), ('Ch', 'Ч'), ('ch', 'ч'),
    ('NG', 'НГ'), ('Ng', 'Нг'), ('ng', 'нг'),
    ('YA', 'Я'), ('Ya', 'Я'), ('ya', 'я'),
    ('YO', 'Ё'), ('Yo', 'Ё'), ('yo', 'ё'),
    ('YU', 'Ю'), ('Yu', 'Ю'), ('yu', 'ю'),
    ('TS', 'Ц'), ('Ts', 'Ц'), ('ts', 'ц'),
]

LATIN_SINGLE = {
    'A': 'А', 'a': 'а', 'B': 'Б', 'b': 'б',
    'D': 'Д', 'd': 'д', 'E': 'Е', 'e': 'е',
    'F': 'Ф', 'f': 'ф', 'G': 'Г', 'g': 'г',
    'H': 'Х', 'h': 'х', 'I': 'И', 'i': 'и',
    'J': 'Ж', 'j': 'ж', 'K': 'К', 'k': 'к',
    'L': 'Л', 'l': 'л', 'M': 'М', 'm': 'м',
    'N': 'Н', 'n': 'н', 'O': 'О', 'o': 'о',
    'P': 'П', 'p': 'п', 'Q': 'К', 'q': 'к',
    'R': 'Р', 'r': 'р', 'S': 'С', 's': 'с',
    'T': 'Т', 't': 'т', 'U': 'У', 'u': 'у',
    'V': 'В', 'v': 'в', 'X': 'Х', 'x': 'х',
    'Y': 'Й', 'y': 'й', 'Z': 'З', 'z': 'з',
    "'": '', '`': '',
}


def is_latin(text):
    latin = sum(1 for c in text if 'A' <= c <= 'Z' or 'a' <= c <= 'z')
    cyrillic = sum(1 for c in text if 'Ѐ' <= c <= 'ӿ')
    return latin > cyrillic


def normalize_name(text):
    if not text or text == 'б/о':
        return text
    if is_latin(text):
        result = text
        for src, dst in LATIN_MULTI:
            result = result.replace(src, dst)
        return ''.join(LATIN_SINGLE.get(c, c) for c in result)
    else:
        result = text
        for src, dst in UZB_CYRILLIC_MAP.items():
            result = result.replace(src, dst)
        return result


def detect_columns(headers):
    """FIO va sana ustunlarini topadi. Aniqroq kalit so'zlar ustunlik qiladi."""
    fio_col = None
    date_col = None

    # Aniqroqdan umumiyga qarab tekshirish (birinchi moslik olinadi)
    fio_priority = [
        ['ф.и.о', 'фио', 'ф.и.ш', 'fish'],          # aniq
        ['фамилия', 'familiya', 'исм-шариф'],          # o'rtacha
        ['исм', 'фам'],                                 # umumiy
    ]
    date_priority = [
        ['год рождения', 'дата рождения', 'туғилган сана', "tug'ilgan sana"],  # aniq
        ['туғилган', "tug'ilgan", 'birth', 'число месяц'],                      # o'rtacha
        ['год', 'yil', 'дата', 'sana'],                                         # umumiy
    ]
    # 'место' so'zi bor ustunlarni sana sifatida qabul qilmaymiz
    exclude_date = ['место', 'joyi', 'manzil']

    for priority_list in fio_priority:
        if fio_col is not None:
            break
        for i, h in enumerate(headers):
            if any(kw in h for kw in priority_list):
                fio_col = i
                break

    for priority_list in date_priority:
        if date_col is not None:
            break
        for i, h in enumerate(headers):
            if any(kw in h for kw in priority_list):
                if not any(ex in h for ex in exclude_date):
                    date_col = i
                    break

    if fio_col is None:
        fio_col = 0
    if date_col is None and len(headers) > 1:
        date_col = 1

    return fio_col, date_col


def extract_year(text):
    if not text:
        return ''
    text = str(text)
    match = re.search(r'\b(19|20)\d{2}\b', text)
    if match:
        return int(match.group())
    return ''


def split_fio(fio):
    if not fio:
        return '', '', ''
    # Ortiqcha probellarni tozalash (boshida, oxirida, o'rtasida)
    cleaned = re.sub(r'\s+', ' ', str(fio)).strip()
    parts = cleaned.split()
    if len(parts) >= 3:
        f, i, s = parts[0], parts[1], ' '.join(parts[2:])
    elif len(parts) == 2:
        f, i, s = parts[0], parts[1], 'б/о'
    elif len(parts) == 1:
        f, i, s = parts[0], '', 'б/о'
    else:
        return '', '', ''
    return normalize_name(f), normalize_name(i), normalize_name(s)


def read_excel_file(filepath):
    """Excel fayldan FIO va tug'ilgan sana ustunlarini topib o'qiydi"""
    wb = openpyxl.load_workbook(filepath, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    # Ustun nomlarini topish
    header_row = None
    header_idx = 0
    for i, row in enumerate(rows[:10]):
        row_text = ' '.join(str(c).lower() for c in row if c)
        if any(w in row_text for w in ['фамилия', 'ф.и.о', 'фио', 'исм', 'familiya', 'fish']):
            header_row = row
            header_idx = i
            break

    if header_row is None:
        # Birinchi qator header deb qabul qilamiz
        header_row = rows[0]
        header_idx = 0

    headers = [str(c).strip().lower() if c else '' for c in header_row]

    # FIO va sana ustunini aniqlash
    fio_col = None
    date_col = None

    fio_col, date_col = detect_columns(headers)

    results = []
    for row in rows[header_idx + 1:]:
        if not any(row):
            continue
        fio = row[fio_col] if fio_col < len(row) else ''
        sana = row[date_col] if date_col is not None and date_col < len(row) else ''
        if fio:
            familiya, ismi, sharif = split_fio(fio)
            yil = extract_year(sana)
            results.append((familiya, ismi, sharif, yil))

    return results


def read_word_file(filepath):
    """Word fayldan FIO va tug'ilgan sana ma'lumotlarini o'qiydi"""
    doc = Document(filepath)
    results = []

    # Jadvallardan o'qish
    for table in doc.tables:
        if not table.rows:
            continue

        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        fio_col = None
        date_col = None

        fio_col, date_col = detect_columns(header_cells)

        # Ikkinchi qator raqamlardan iborat bo'lsa (1,2,3...) uni o'tkazib yuboramiz
        data_start = 1
        if len(table.rows) > 1:
            second_row = [c.text.strip() for c in table.rows[1].cells]
            if all(c.isdigit() or c == '' for c in second_row if c):
                data_start = 2

        for row in table.rows[data_start:]:
            cells = [c.text.strip() for c in row.cells]
            fio = cells[fio_col] if fio_col < len(cells) else ''
            sana = cells[date_col] if date_col is not None and date_col < len(cells) else ''
            if fio:
                familiya, ismi, sharif = split_fio(fio)
                yil = extract_year(sana)
                results.append((familiya, ismi, sharif, yil))

    # Jadval topilmasa matndan o'qish
    if not results:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Raqam bilan boshlangan qatorlar (ro'yxat)
            match = re.match(r'^\d+[\.\)]\s+(.+)', text)
            if match:
                content = match.group(1)
                parts = content.split(',')
                fio = parts[0].strip()
                sana = parts[1].strip() if len(parts) > 1 else ''
                familiya, ismi, sharif = split_fio(fio)
                yil = extract_year(sana)
                if familiya:
                    results.append((familiya, ismi, sharif, yil))

    return results


def create_output_excel(data, output_path):
    """Template formatida Excel fayl yaratish"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    # Sarlavha uslubi
    header_font = Font(name='Times New Roman', bold=True, size=11)
    cell_font = Font(name='Times New Roman', size=11)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    thin = Side(style='thin')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Sarlavhalar
    headers = ['п/п', 'Фамилия', 'Имя', 'Отчество', 'Год рождения']
    ws.append(headers)

    for col_idx, _ in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    # Ma'lumotlar
    for i, (familiya, ismi, sharif, yil) in enumerate(data, 1):
        ws.append([i, familiya, ismi, sharif, yil])
        for col_idx in range(1, 6):
            cell = ws.cell(row=i + 1, column=col_idx)
            cell.font = cell_font
            cell.border = border
            cell.alignment = center if col_idx in (1, 5) else left

    # Ustun kengliklari
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 22
    ws.column_dimensions['E'].width = 15

    ws.row_dimensions[1].height = 30

    wb.save(output_path)


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("FIO Ajratuvchi Dastur")
        self.geometry("600x480")
        self.resizable(False, False)
        self.configure(bg='#f0f0f0')

        self.files = []
        self.build_ui()

    def build_ui(self):
        title = tk.Label(self, text="FIO va Tug'ilgan Yil Ajratuvchi",
                         font=('Arial', 14, 'bold'), bg='#f0f0f0', fg='#333')
        title.pack(pady=15)

        # Fayl qo'shish
        btn_frame = tk.Frame(self, bg='#f0f0f0')
        btn_frame.pack(pady=5)

        tk.Button(btn_frame, text="+ Fayl qo'shish (Excel/Word)",
                  command=self.add_files, bg='#4CAF50', fg='white',
                  font=('Arial', 11), padx=10, pady=5).pack(side='left', padx=5)

        tk.Button(btn_frame, text="Tozalash",
                  command=self.clear_files, bg='#f44336', fg='white',
                  font=('Arial', 11), padx=10, pady=5).pack(side='left', padx=5)

        # Fayl ro'yxati
        list_frame = tk.Frame(self, bg='#f0f0f0')
        list_frame.pack(fill='both', padx=20, pady=5)

        tk.Label(list_frame, text="Tanlangan fayllar:", font=('Arial', 10),
                 bg='#f0f0f0', anchor='w').pack(fill='x')

        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side='right', fill='y')

        self.file_listbox = tk.Listbox(list_frame, height=8, yscrollcommand=scrollbar.set,
                                        font=('Arial', 9), selectmode='single')
        self.file_listbox.pack(fill='both')
        scrollbar.config(command=self.file_listbox.yview)

        # Progress
        self.progress = ttk.Progressbar(self, length=560, mode='determinate')
        self.progress.pack(pady=10, padx=20)

        self.status_label = tk.Label(self, text="Fayl tanlang va ishga tushiring",
                                      font=('Arial', 10), bg='#f0f0f0', fg='#555')
        self.status_label.pack()

        # Ishga tushirish
        tk.Button(self, text="Ishga tushirish",
                  command=self.run, bg='#2196F3', fg='white',
                  font=('Arial', 12, 'bold'), padx=20, pady=8).pack(pady=15)

    def add_files(self):
        files = filedialog.askopenfilenames(
            title="Fayllarni tanlang",
            filetypes=[("Excel va Word fayllar", "*.xlsx *.xls *.docx"), ("Barchasi", "*.*")]
        )
        for f in files:
            if f not in self.files:
                self.files.append(f)
                self.file_listbox.insert('end', os.path.basename(f))

    def clear_files(self):
        self.files.clear()
        self.file_listbox.delete(0, 'end')
        self.status_label.config(text="Fayl tanlang va ishga tushiring")

    def run(self):
        if not self.files:
            messagebox.showwarning("Ogohlantirish", "Avval fayl tanlang!")
            return

        output_path = filedialog.asksaveasfilename(
            title="Natijani saqlash",
            defaultextension=".xlsx",
            filetypes=[("Excel fayl", "*.xlsx")],
            initialfile="natija.xlsx"
        )
        if not output_path:
            return

        all_data = []
        total = len(self.files)
        self.progress['maximum'] = total

        for i, filepath in enumerate(self.files):
            self.status_label.config(text=f"O'qilmoqda: {os.path.basename(filepath)}")
            self.update()
            try:
                ext = os.path.splitext(filepath)[1].lower()
                if ext in ('.xlsx', '.xls'):
                    data = read_excel_file(filepath)
                elif ext == '.docx':
                    data = read_word_file(filepath)
                else:
                    continue
                all_data.extend(data)
            except Exception as e:
                messagebox.showerror("Xato", f"{os.path.basename(filepath)}: {e}")
            self.progress['value'] = i + 1
            self.update()

        if not all_data:
            messagebox.showwarning("Natija", "Hech qanday ma'lumot topilmadi!")
            return

        create_output_excel(all_data, output_path)
        self.status_label.config(text=f"Tayyor! {len(all_data)} ta yozuv saqlandi.")
        self.progress['value'] = 0

        if messagebox.askyesno("Muvaffaqiyat",
                                f"{len(all_data)} ta yozuv saqlandi.\nFaylni ochishni xohlaysizmi?"):
            os.system(f'open "{output_path}"')


if __name__ == '__main__':
    app = App()
    app.mainloop()
